import os
import docx
from docx.opc.constants import RELATIONSHIP_TYPE
import logging

def extract_docx(file_path: str, document_id: str, images_dir: str):
    """
    Extracts text, headings, and tables from a DOCX file.
    """
    extracted_data = {
        "pages": [], # DOCX doesn't have strict pages, we'll map everything to page 1 or chunk by paragraphs
        "headings": [],
        "images": [],
        "diagrams": [],
        "tables": []
    }

    try:
        doc = docx.Document(file_path)
        
        text_content = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            text_content.append(text)
            
            # Detect Headings
            if para.style.name.startswith('Heading'):
                level_str = para.style.name.replace('Heading ', '')
                level = int(level_str) if level_str.isdigit() else 1
                extracted_data["headings"].append({
                    "page_number": 1,
                    "level": level,
                    "text": text
                })
        
        # Save aggregated text
        extracted_data["pages"].append({
            "page_number": 1,
            "text_content": "\n".join(text_content)
        })
        
        # Tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
                
            if table_data:
                headers = table_data[0]
                rows = table_data[1:] if len(table_data) > 1 else []
                extracted_data["tables"].append({
                    "page_number": 1,
                    "headers": headers,
                    "rows": rows
                })

        # Image extraction is more complex in python-docx, typically handled by parsing rels
        for rel in doc.part.rels.values():
            if rel.reltype == RELATIONSHIP_TYPE.IMAGE:
                try:
                    img_data = rel.target_part.blob
                    img_ext = rel.target_part.content_type.split('/')[-1]
                    img_filename = f"{document_id}_img_{len(extracted_data['images']) + 1}.{img_ext}"
                    img_path = os.path.join(images_dir, img_filename)
                    
                    with open(img_path, "wb") as f:
                        f.write(img_data)
                        
                    extracted_data["images"].append({
                        "page_number": 1,
                        "image_path": img_path
                    })
                except Exception as e:
                    logging.warning(f"Failed to extract DOCX image: {str(e)}")

    except Exception as e:
        logging.error(f"Error processing DOCX {file_path}: {str(e)}")
        raise e

    return extracted_data
